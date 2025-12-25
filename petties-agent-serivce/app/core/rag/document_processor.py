"""
PETTIES AGENT SERVICE - Document Processor

Uses LlamaIndex for document parsing, chunking.
Uses Cohere for multilingual embeddings (Vietnamese support).

Package: app.core.rag
Purpose: Document processing pipeline for RAG
Version: v1.0.0 (Migrated to Cohere embeddings)

Changes from v0.x:
- Replaced OpenAI embeddings with Cohere embed-multilingual-v3.0
- Better Vietnamese text support
- Async embedding with batching
"""

from llama_index.core import Document, Settings
from llama_index.core.node_parser import SentenceSplitter
from typing import List, Optional
from loguru import logger
import io

from app.config.settings import settings


class DocumentProcessor:
    """
    Document processing pipeline using LlamaIndex + Cohere

    Usage:
        processor = DocumentProcessor()
        chunks = processor.process_file(file_bytes, "doc.pdf")
        embeddings = await processor.embed_chunks(chunks)
    """

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
    ):
        """
        Initialize Document Processor

        Args:
            chunk_size: Maximum chunk size in characters
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self._embedding_client = None

    async def _get_embedding_client(self):
        """Get Cohere embedding client (lazy initialization)"""
        if self._embedding_client is None:
            from app.services.embeddings import create_embedding_client, EmbeddingConfig

            self._embedding_client = create_embedding_client(EmbeddingConfig(
                provider="cohere",
                model="embed-multilingual-v3.0",
                api_key=settings.COHERE_API_KEY
            ))

        return self._embedding_client

    def process_file(
        self,
        file_content: bytes,
        filename: str,
        metadata: Optional[dict] = None
    ) -> List[dict]:
        """
        Process file into chunks

        Args:
            file_content: Raw file bytes
            filename: Original filename
            metadata: Additional metadata

        Returns:
            List of chunk dicts with text and metadata
        """
        # Detect file type
        ext = filename.lower().split('.')[-1]

        text = self._extract_text(file_content, ext)
        if not text:
            logger.warning(f"No text extracted from {filename}")
            return []

        # Create LlamaIndex document
        doc = Document(
            text=text,
            metadata={
                "filename": filename,
                "file_type": ext,
                **(metadata or {})
            }
        )

        # Split into nodes/chunks
        nodes = self.splitter.get_nodes_from_documents([doc])

        chunks = []
        for i, node in enumerate(nodes):
            chunks.append({
                "chunk_index": i,
                "content": node.text,
                "metadata": {
                    **node.metadata,
                    "chunk_index": i,
                    "total_chunks": len(nodes)
                }
            })

        logger.info(f"Processed {filename}: {len(chunks)} chunks")
        return chunks

    def _extract_text(self, content: bytes, ext: str) -> str:
        """Extract text from file based on extension"""
        try:
            if ext == "txt" or ext == "md":
                return content.decode("utf-8")

            elif ext == "pdf":
                try:
                    # Try PyMuPDF first (better quality)
                    import fitz
                    doc = fitz.open(stream=content, filetype="pdf")
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    return text
                except ImportError:
                    logger.warning("PyMuPDF not installed, trying PyPDF2")
                    try:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(io.BytesIO(content))
                        return "\n".join(page.extract_text() or "" for page in reader.pages)
                    except ImportError:
                        logger.error("No PDF library available. Install PyMuPDF or PyPDF2")
                        return ""

            elif ext in ["doc", "docx"]:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(content))
                return "\n".join(p.text for p in doc.paragraphs)

            else:
                # Try as plain text
                return content.decode("utf-8", errors="ignore")

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return ""

    async def embed_chunks(self, chunks: List[dict]) -> List[List[float]]:
        """
        Generate embeddings for chunks using Cohere

        Args:
            chunks: List of chunk dicts with 'content' key

        Returns:
            List of embedding vectors (1024 dimensions)
        """
        if not chunks:
            return []

        texts = [c["content"] for c in chunks]

        # Get Cohere embedding client
        embedding_client = await self._get_embedding_client()

        try:
            result = await embedding_client.embed_documents(texts)
            logger.info(f"Generated {len(result.embeddings)} embeddings with Cohere")
            return result.embeddings

        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise

    async def embed_query(self, query: str) -> List[float]:
        """
        Generate embedding for a query using Cohere

        Args:
            query: Query text

        Returns:
            Embedding vector (1024 dimensions)
        """
        embedding_client = await self._get_embedding_client()

        try:
            return await embedding_client.embed_query(query)

        except Exception as e:
            logger.error(f"Query embedding failed: {e}")
            raise


# Singleton instance
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get singleton DocumentProcessor instance"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP
        )
    return _processor
