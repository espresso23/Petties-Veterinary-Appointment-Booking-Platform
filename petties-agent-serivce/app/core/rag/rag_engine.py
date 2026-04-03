"""
PETTIES AGENT SERVICE - RAG Engine (Full LlamaIndex)

Full LlamaIndex integration for document indexing and retrieval.
Uses LlamaIndex to manage: chunking, embedding, vector storage, and search.

Package: app.core.rag
Purpose: Document indexing and retrieval for pet care knowledge
Version: v2.0.0 (Full LlamaIndex integration)

Components:
- LlamaIndex VectorStoreIndex for indexing
- Cohere embed-multilingual-v3.0 for Vietnamese embeddings
- Qdrant Cloud for vector storage
"""

from typing import List, Optional, Any, Tuple
from dataclasses import dataclass
from loguru import logger
import asyncio
import base64
import io
from pathlib import Path
import hashlib

# LlamaIndex imports
from llama_index.core import (
    VectorStoreIndex,
    Document,
    Settings,
    StorageContext,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.cohere import CohereEmbedding
# from llama_index.llms.openrouter import OpenRouter - Moved to initialize
# from llama_index.vector_stores.qdrant import QdrantVectorStore - Moved to initialize
# from qdrant_client import QdrantClient - Moved to initialize

from app.config.settings import settings
from app.core.config_helper import get_setting
from app.db.postgres.session import AsyncSessionLocal


# Cohere embed-multilingual-v3.0 dimension
COHERE_EMBED_DIMENSION = 1024

# Knowledge Base Image Collection (for image vectors extracted from PDFs)
KB_IMAGE_COLLECTION = "petties_kb_images"
JINA_IMAGE_DIMENSION = 1024

# Image extraction settings
MAX_IMAGES_PER_DOCUMENT = 50  # Tăng từ 10 để hỗ trợ PDF lớn (300+ trang)
IMAGE_MIN_SIZE_BYTES = 1000  # Skip tiny images likely to be logos/icons


@dataclass
class RetrievedChunk:
    """Retrieved document chunk"""

    document_id: int
    document_name: str
    chunk_index: int
    content: str
    score: float


@dataclass
class IndexResult:
    """Result from indexing a document (text + images)"""

    text_chunks: int
    image_vectors: int


class LlamaIndexRAGEngine:
    """
    Full LlamaIndex RAG Engine

    Uses LlamaIndex to handle:
    - Document chunking (SentenceSplitter)
    - Embedding (Cohere)
    - Vector storage (Qdrant)
    - Similarity search

    Usage:
        engine = LlamaIndexRAGEngine()
        await engine.initialize()
        await engine.index_document(content, "doc.pdf", doc_id=1)
        results = await engine.query("pet symptoms")
    """

    _initialized: bool = False

    def __init__(self):
        if LlamaIndexRAGEngine._initialized:
            return

        self.index: Optional[VectorStoreIndex] = None
        self.vector_store: Optional[Any] = None
        self.qdrant_client: Optional[Any] = None
        self._collection_name = (
            settings.QDRANT_COLLECTION_NAME or "petties_knowledge_base"
        )
        # Image vector support
        self._image_qdrant_client: Optional[Any] = None
        self._image_collection_initialized: bool = False
        self._init_lock = asyncio.Lock()
        LlamaIndexRAGEngine._initialized = True

    async def initialize(self):
        """
        Initialize LlamaIndex components with settings from database

        Must be called before using index_document or query
        """
        if self.index is not None:
            return

        async with self._init_lock:
            if self.index is not None:
                return

            logger.info("Initializing LlamaIndex RAG Engine...")

            # Lazy imports
            from qdrant_client import QdrantClient
            from llama_index.vector_stores.qdrant import QdrantVectorStore
            from llama_index.llms.openrouter import OpenRouter

            async with AsyncSessionLocal() as db:
                # Get API keys from database
                cohere_api_key = await get_setting("COHERE_API_KEY", db)
                cohere_model = (
                    await get_setting("COHERE_EMBEDDING_MODEL", db)
                    or "embed-multilingual-v3.0"
                )
                qdrant_url = await get_setting("QDRANT_URL", db) or settings.QDRANT_URL
                qdrant_api_key = (
                    await get_setting("QDRANT_API_KEY", db) or settings.QDRANT_API_KEY
                )
                self._collection_name = (
                    await get_setting("QDRANT_COLLECTION_NAME", db)
                    or "petties_knowledge_base"
                )
                openrouter_api_key = await get_setting("OPENROUTER_API_KEY", db)
                llm_model = (
                    await get_setting("RAG_LLM_MODEL", db)
                    or "google/gemini-2.5-flash-lite"
                )

            if not cohere_api_key:
                logger.warning(
                    "COHERE_API_KEY not configured. RAG search will be unavailable. Please set it in Settings."
                )
                return

            # Configure LlamaIndex Settings (global)
            Settings.embed_model = CohereEmbedding(
                api_key=cohere_api_key,
                model_name=cohere_model,
                input_type="search_document",  # For indexing
            )

            # Configure LLM (OpenRouter)
            if openrouter_api_key:
                Settings.llm = OpenRouter(
                    api_key=openrouter_api_key,
                    model=llm_model,
                    temperature=0.1,
                )
            else:
                logger.warning(
                    "OPENROUTER_API_KEY not configured. RAG synthesis may fail if LLM is needed."
                )

            # Configure chunking
            Settings.node_parser = SentenceSplitter(chunk_size=512, chunk_overlap=50)

            # Initialize Qdrant client
            if qdrant_url and qdrant_api_key:
                logger.info(f"Connecting to Qdrant Cloud: {qdrant_url}")
                self.qdrant_client = QdrantClient(
                    url=qdrant_url, api_key=qdrant_api_key
                )
            else:
                logger.info("Using local Qdrant")
                self.qdrant_client = QdrantClient(host="localhost", port=6333)

            # Create vector store
            self.vector_store = QdrantVectorStore(
                client=self.qdrant_client,
                collection_name=self._collection_name,
                enable_hybrid=False,  # Can enable for BM25 + Vector
            )

            # Create or load index
            storage_context = StorageContext.from_defaults(
                vector_store=self.vector_store
            )

            # Check if collection exists with data
            try:
                collection_info = self.qdrant_client.get_collection(
                    self._collection_name
                )
                points_count = collection_info.points_count

                if points_count > 0:
                    logger.info(f"Loading existing index with {points_count} vectors")
                    self.index = VectorStoreIndex.from_vector_store(
                        self.vector_store, storage_context=storage_context
                    )
                else:
                    logger.info("Creating new empty index")
                    self.index = VectorStoreIndex.from_documents(
                        [], storage_context=storage_context
                    )
            except Exception as e:
                logger.warning(f"Collection not found, creating new: {e}")
                self.index = VectorStoreIndex.from_documents(
                    [], storage_context=storage_context
                )

            self._initialized = True
            logger.info("LlamaIndex RAG Engine initialized successfully")

    async def index_document(
        self,
        file_content: Optional[bytes] = None,
        filename: Optional[str] = None,
        document_id: Optional[int] = None,
        file_path: Optional[Path] = None,
        metadata: Optional[dict] = None,
    ) -> IndexResult:
        """
        Index a document into the knowledge base

        LlamaIndex handles:
        - Text extraction (if needed)
        - Chunking with SentenceSplitter
        - Embedding with Cohere
        - Storing in Qdrant

        Args:
            file_content: Raw file bytes (optional if file_path is provided)
            filename: Original filename
            document_id: Database document ID
            file_path: Path to file on disk (optional if file_content is provided)
            metadata: Additional metadata

        Returns:
            Number of chunks indexed
        """
        await self.initialize()

        # Extract text from file (non-blocking)
        if file_path and file_path.exists():
            text = await asyncio.to_thread(self._extract_text_from_path, file_path)
            if not filename:
                filename = file_path.name
        elif file_content:
            text = await asyncio.to_thread(
                self._extract_text, file_content, filename or "unknown"
            )
        else:
            logger.error("No file content or path provided for indexing")
            return 0

        if not text:
            logger.warning(f"No text extracted from {filename}")
            return IndexResult(text_chunks=0, image_vectors=0)

        # Create LlamaIndex Document with metadata
        doc_metadata = {
            "document_id": document_id,
            "document_name": filename,
            "filename": filename,
            "file_type": filename.split(".")[-1].lower(),
            **(metadata or {}),
        }

        doc = Document(text=text, metadata=doc_metadata, doc_id=str(document_id))

        # Insert into index (LlamaIndex handles chunking + embedding + storage)
        try:
            # Delete old chunks for this document first to avoid duplicates
            if document_id is not None:
                try:
                    deleted = await self.delete_document(document_id)
                    if deleted:
                        logger.info(
                            f"Removed old vectors for document_id={document_id} before re-index"
                        )
                except Exception as del_err:
                    logger.warning(
                        f"Could not delete old vectors (non-fatal): {del_err}"
                    )

            # Insert new document
            self.index.insert(doc)

            # Count nodes created
            nodes = Settings.node_parser.get_nodes_from_documents([doc])
            chunks_count = len(nodes)

            logger.info(f"Indexed {filename}: {chunks_count} chunks with LlamaIndex")

            # Index images from PDF (if applicable)
            images_indexed = 0
            if file_path and file_path.exists() and file_path.suffix.lower() == ".pdf":
                try:
                    images_indexed = await self._index_document_images(
                        file_path=file_path,
                        document_id=document_id,
                        filename=filename,
                        metadata=metadata,
                    )
                    if images_indexed > 0:
                        logger.info(f"Also indexed {images_indexed} images from PDF")
                except Exception as img_err:
                    logger.warning(f"Image indexing failed (non-fatal): {img_err}")

            return IndexResult(text_chunks=chunks_count, image_vectors=images_indexed)

        except Exception as e:
            logger.error(f"Failed to index document: {e}")
            raise

    async def query(
        self,
        query: str,
        top_k: int = 5,
        min_score: float = 0.5,
        document_ids: Optional[List[int]] = None,
    ) -> List[RetrievedChunk]:
        """
        Query the knowledge base

        LlamaIndex handles:
        - Query embedding with Cohere
        - Vector similarity search in Qdrant

        Args:
            query: Search query text
            top_k: Number of results
            min_score: Minimum similarity score
            document_ids: Filter by specific documents (optional)

        Returns:
            List of retrieved chunks
        """
        await self.initialize()

        logger.info(f"Query: '{query[:50]}...', top_k={top_k}, min_score={min_score}")

        try:
            # Create retriever with settings
            retriever = self.index.as_retriever(
                similarity_top_k=top_k,
            )

            # Retrieve nodes
            nodes = await asyncio.to_thread(retriever.retrieve, query)

            logger.info(f"Retrieved {len(nodes)} raw results")

            # Filter by score, deduplicate, and convert to RetrievedChunk
            chunks = []
            seen_contents: set[str] = set()
            for i, node in enumerate(nodes):
                score = node.score if hasattr(node, "score") and node.score else 0.0

                # Skip if below min_score
                if score < min_score:
                    continue

                # Extract metadata
                meta = node.metadata if hasattr(node, "metadata") else {}
                content = node.text if hasattr(node, "text") else str(node)

                # Skip duplicates
                content_key = content.strip()
                if content_key in seen_contents:
                    logger.debug(
                        f"  Skipping duplicate chunk (doc='{meta.get('document_name', '?')}', score={score:.3f})"
                    )
                    continue
                seen_contents.add(content_key)

                chunk = RetrievedChunk(
                    document_id=meta.get("document_id", 0),
                    document_name=meta.get("document_name", meta.get("filename", "")),
                    chunk_index=i,
                    content=content,
                    score=score,
                )
                chunks.append(chunk)

                logger.debug(
                    f"  Result {len(chunks) - 1}: score={score:.3f}, content_len={len(content)}, doc='{chunk.document_name}'"
                )

            # Filter by document_ids if provided
            if document_ids:
                chunks = [c for c in chunks if c.document_id in document_ids]

            doc_names = [c.document_name for c in chunks]
            logger.info(
                f"Query '{query[:30]}...' returned {len(chunks)} chunks from docs: {doc_names}"
            )
            return chunks

        except Exception as e:
            logger.error(f"Query failed: {e}")
            import traceback

            logger.error(traceback.format_exc())
            return []

    async def delete_document(self, document_id: int) -> int:
        """
        Delete all chunks for a document from Qdrant

        Args:
            document_id: Database document ID

        Returns:
            Number of points deleted
        """
        await self.initialize()

        try:
            # Delete by filter
            from qdrant_client.models import Filter, FieldCondition, MatchValue

            result = self.qdrant_client.delete(
                collection_name=self._collection_name,
                points_selector=Filter(
                    must=[
                        FieldCondition(
                            key="document_id", match=MatchValue(value=document_id)
                        )
                    ]
                ),
            )

            logger.info(f"Deleted vectors for document_id={document_id}")
            return 1  # Qdrant doesn't return count

        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            return 0

    def _extract_text_from_path(self, path: Path) -> str:
        """Extract text from file path safely"""
        try:
            ext = path.suffix.lower()[1:]
            if ext in ["txt", "md"]:
                return path.read_text(encoding="utf-8", errors="replace")

            with open(path, "rb") as f:
                content = f.read()
                return self._extract_text(content, path.name)
        except Exception as e:
            logger.error(f"Text extraction failed for {path}: {e}")
            return ""

    def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from file based on extension"""
        ext = filename.lower().split(".")[-1]

        try:
            if ext in ["txt", "md"]:
                return content.decode("utf-8")

            elif ext == "pdf":
                try:
                    import fitz  # PyMuPDF

                    doc = fitz.open(stream=content, filetype="pdf")
                    text = ""
                    for page in doc:
                        text += page.get_text()
                    return text
                except ImportError:
                    from PyPDF2 import PdfReader
                    import io

                    reader = PdfReader(io.BytesIO(content))
                    return "\n".join(page.extract_text() or "" for page in reader.pages)

            elif ext in ["doc", "docx"]:
                from docx import Document as DocxDocument
                import io

                doc = DocxDocument(io.BytesIO(content))
                return "\n".join(p.text for p in doc.paragraphs)

            else:
                return content.decode("utf-8", errors="ignore")

        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return ""

    def _extract_images_from_pdf(self, file_path: Path) -> List[Tuple[str, bytes, str]]:
        """
        Extract images from PDF using PyMuPDF (fitz).

        Returns:
            List of (image_id, image_bytes, extension) tuples
        """
        images = []
        try:
            import fitz

            with fitz.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf):
                    image_list = page.get_images(full=True)
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = pdf.extract_image(xref)
                            image_bytes = base_image["image"]
                            image_ext = base_image.get("ext", "jpeg").lower()
                            # Skip tiny images (likely logos/icons)
                            if len(image_bytes) < IMAGE_MIN_SIZE_BYTES:
                                continue
                            image_id = f"p{page_num + 1}_img{img_index + 1}_{hashlib.md5(image_bytes[:100]).hexdigest()[:8]}"
                            images.append((image_id, image_bytes, image_ext))
                            if len(images) >= MAX_IMAGES_PER_DOCUMENT:
                                logger.warning(
                                    f"Reached max images per document ({MAX_IMAGES_PER_DOCUMENT})"
                                )
                                break
                        except Exception as e:
                            logger.debug(
                                f"Failed to extract image {img_index} from page {page_num}: {e}"
                            )
                            continue
                    if len(images) >= MAX_IMAGES_PER_DOCUMENT:
                        break
            logger.info(f"Extracted {len(images)} images from PDF")
        except ImportError:
            logger.warning(
                "PyMuPDF (fitz) not installed - cannot extract images from PDF"
            )
        except Exception as e:
            logger.error(f"Error extracting images from PDF: {e}")
        return images

    async def _ensure_image_collection(self):
        """Ensure KB image collection exists with named vectors (text + image)."""
        if self._image_collection_initialized:
            return

        try:
            from qdrant_client import QdrantClient
            from qdrant_client.models import VectorParams, Distance

            # Get Qdrant config from DB settings (same as text RAG)
            async with AsyncSessionLocal() as db:
                qdrant_url = await get_setting("QDRANT_URL", db)
                qdrant_api_key = await get_setting("QDRANT_API_KEY", db)

            if qdrant_url and qdrant_api_key:
                self._image_qdrant_client = QdrantClient(
                    url=qdrant_url, api_key=qdrant_api_key
                )
            else:
                logger.warning(
                    "QDRANT_URL/QDRANT_API_KEY not in DB settings - image indexing disabled"
                )
                return

            # Create collection if not exists
            try:
                exists = self._image_qdrant_client.collection_exists(
                    KB_IMAGE_COLLECTION
                )
            except Exception:
                exists = False

            if not exists:
                self._image_qdrant_client.create_collection(
                    collection_name=KB_IMAGE_COLLECTION,
                    vectors_config={
                        "text": VectorParams(
                            size=COHERE_EMBED_DIMENSION,
                            distance=Distance.COSINE,
                        ),
                        "image": VectorParams(
                            size=JINA_IMAGE_DIMENSION,
                            distance=Distance.COSINE,
                        ),
                    },
                )
                logger.info(f"Created KB image collection: {KB_IMAGE_COLLECTION}")

            self._image_collection_initialized = True
            logger.info("KB image collection initialized")

        except Exception as e:
            logger.error(f"Failed to initialize image collection: {e}")
            self._image_collection_initialized = False

    async def _index_document_images(
        self,
        file_path: Path,
        document_id: int,
        filename: str,
        metadata: Optional[dict] = None,
    ) -> int:
        """
        Extract images from PDF and index them into KB image collection.

        Returns:
            Number of images indexed
        """
        if file_path.suffix.lower() != ".pdf":
            return 0

        await self._ensure_image_collection()

        if not self._image_qdrant_client:
            logger.debug("Image Qdrant client not available - skipping image indexing")
            return 0

        images = self._extract_images_from_pdf(file_path)
        if not images:
            logger.info(f"No images extracted from PDF {filename}")
            return 0

        try:
            from app.core.embeddings.jina_image_embeddings import embed_image_base64

            # Convert images to base64 for Jina
            base64_images = []
            for img_id, img_bytes, img_ext in images:
                b64 = base64.b64encode(img_bytes).decode("utf-8")
                mime_type = f"image/{img_ext}" if img_ext else "image/jpeg"
                base64_images.append(f"data:{mime_type};base64,{b64}")

            # Get image embeddings
            embeddings = await embed_image_base64(base64_images)

            if not embeddings:
                logger.warning("No image embeddings generated")
                return 0

            # Get text for text vectors (use document name as fallback)
            text_for_embedding = metadata.get("notes", "") if metadata else ""
            if not text_for_embedding:
                text_for_embedding = f"Document: {filename}"

            # Create text embeddings (reuse existing Cohere)
            await self.initialize()
            if self.index:
                # Use Cohere embedding for text
                from llama_index.embeddings.cohere import CohereEmbedding

                embed_model = CohereEmbedding(
                    model="embed-multilingual-v3.0",
                    api_key=(await self._get_cohere_api_key()),
                )
                text_embedding = await asyncio.to_thread(
                    embed_model.get_text_embedding, text_for_embedding
                )
            else:
                text_embedding = None

            # Upsert points to Qdrant with named vectors
            from qdrant_client.models import PointStruct
            from datetime import datetime

            points = []
            for i, (img_id, img_bytes, img_ext) in enumerate(images[: len(embeddings)]):
                point = PointStruct(
                    id=f"{document_id}_{img_id}",
                    vector={
                        "text": text_embedding
                        if text_embedding
                        else [0.0] * COHERE_EMBED_DIMENSION,
                        "image": embeddings[i],
                    },
                    payload={
                        "document_id": document_id,
                        "filename": filename,
                        "image_id": img_id,
                        "image_index": i,
                        "extracted_at": datetime.utcnow().isoformat(),
                        "metadata": metadata or {},
                    },
                )
                points.append(point)

            self._image_qdrant_client.upsert(
                collection_name=KB_IMAGE_COLLECTION,
                points=points,
            )

            logger.info(f"Indexed {len(points)} images for document {document_id}")
            return len(points)

        except Exception as e:
            logger.error(f"Failed to index document images: {e}")
            return 0

    async def _get_cohere_api_key(self) -> Optional[str]:
        """Get Cohere API key from settings."""
        try:
            async with AsyncSessionLocal() as db:
                return await get_setting("COHERE_API_KEY", db)
        except Exception:
            return None

    async def query_with_images(
        self,
        query: str,
        image_urls: Optional[List[str]] = None,
        top_k: int = 5,
        min_score: float = 0.5,
        document_ids: Optional[List[int]] = None,
    ) -> dict:
        """
        Query the knowledge base with text and/or images (hybrid search).

        Args:
            query: Search query text
            image_urls: Optional image URLs to search by image similarity
            top_k: Number of results
            min_score: Minimum similarity score
            document_ids: Filter by specific documents (optional)

        Returns:
            Dict with text_results, image_results, and combined results
        """
        await self.initialize()

        text_results = await self.query(query, top_k, min_score, document_ids)

        image_results = []
        if image_urls and self._image_qdrant_client:
            try:
                from app.core.embeddings.jina_image_embeddings import embed_image_urls
                from qdrant_client.models import Filter, FieldCondition, MatchValue

                image_embeddings = await embed_image_urls(image_urls[:1])
                if image_embeddings:
                    # Query image vector
                    search_result = self._image_qdrant_client.search(
                        collection_name=KB_IMAGE_COLLECTION,
                        query_vector=("image", image_embeddings[0]),
                        limit=top_k,
                        score_threshold=min_score,
                    )

                    for hit in search_result:
                        image_results.append(
                            {
                                "document_id": hit.payload.get("document_id"),
                                "filename": hit.payload.get("filename"),
                                "image_id": hit.payload.get("image_id"),
                                "score": hit.score,
                                "payload": hit.payload,
                            }
                        )
            except Exception as e:
                logger.error(f"Image search failed: {e}")

        # Combine results
        combined = {
            "text_results": text_results,
            "image_results": image_results,
            "has_image_query": len(image_results) > 0,
            "query": query,
        }

        return combined

    async def get_status(self) -> dict:
        """Get RAG engine status"""
        await self.initialize()

        try:
            collection_info = self.qdrant_client.get_collection(self._collection_name)
            return {
                "initialized": self._initialized,
                "collection_name": self._collection_name,
                "points_count": collection_info.points_count,
                "status": str(collection_info.status),
                "engine": "LlamaIndex",
            }
        except Exception as e:
            return {
                "initialized": self._initialized,
                "error": str(e),
                "engine": "LlamaIndex",
            }

    async def recreate_collection(self) -> bool:
        """Delete and recreate the Qdrant collection"""
        try:
            # Delete existing
            try:
                self.qdrant_client.delete_collection(self._collection_name)
                logger.info(f"Deleted collection: {self._collection_name}")
            except Exception:
                pass

            # Recreate with new dimensions
            from qdrant_client.models import Distance, VectorParams

            self.qdrant_client.create_collection(
                collection_name=self._collection_name,
                vectors_config=VectorParams(
                    size=COHERE_EMBED_DIMENSION, distance=Distance.COSINE
                ),
            )

            # Reinitialize vector store and index
            self._initialized = False
            await self.initialize()

            logger.info(f"Recreated collection: {self._collection_name}")
            return True

        except Exception as e:
            logger.error(f"Failed to recreate collection: {e}")
            return False

    async def get_debug_info(self) -> dict:
        """
        Get detailed debug info including sample points

        Used by /debug/qdrant endpoint for troubleshooting

        Returns:
            Dict with collection info, vector count, and sample points
        """
        await self.initialize()

        try:
            # Get collection info
            info = self.qdrant_client.get_collection(self._collection_name)

            # Get sample points
            results = self.qdrant_client.scroll(
                collection_name=self._collection_name,
                limit=3,
                with_vectors=True,
                with_payload=True,
            )

            sample_points = []
            if results and results[0]:
                for point in results[0]:
                    vector_data = point.vector
                    # Handle both dict and list vector formats
                    if isinstance(vector_data, dict):
                        vector_preview = (
                            list(vector_data.values())[0][:5] if vector_data else None
                        )
                    elif isinstance(vector_data, list):
                        vector_preview = vector_data[:5]
                    else:
                        vector_preview = None

                    sample_points.append(
                        {
                            "id": str(point.id),
                            "payload": point.payload,
                            "vector_preview": vector_preview,
                        }
                    )

            return {
                "collection_name": self._collection_name,
                "vectors_count": info.points_count,
                "status": info.status.value if info.status else "unknown",
                "indexed_vectors_count": info.indexed_vectors_count
                if hasattr(info, "indexed_vectors_count")
                else None,
                "sample_points": sample_points,
                "engine": "LlamaIndex",
            }
        except Exception as e:
            logger.error(f"Debug info failed: {e}")
            return {
                "error": str(e),
                "collection_name": self._collection_name,
                "engine": "LlamaIndex",
            }


# Singleton instance
_rag_engine: Optional[LlamaIndexRAGEngine] = None


def get_rag_engine() -> LlamaIndexRAGEngine:
    """Get singleton RAG engine instance"""
    global _rag_engine
    if _rag_engine is None:
        _rag_engine = LlamaIndexRAGEngine()
    return _rag_engine


def reset_rag_engine():
    """Reset singleton RAG engine"""
    global _rag_engine
    _rag_engine = None


# Backward compatibility alias
RAGEngine = LlamaIndexRAGEngine


# Export for compatibility with existing code
__all__ = [
    "LlamaIndexRAGEngine",
    "RAGEngine",  # Alias for backward compatibility
    "RetrievedChunk",
    "get_rag_engine",
    "reset_rag_engine",
    "COHERE_EMBED_DIMENSION",
]
