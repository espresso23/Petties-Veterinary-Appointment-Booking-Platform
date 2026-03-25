"""
PETTIES AGENT SERVICE - Knowledge Base API Routes
REST API endpoints for Document Upload and RAG Query (KB-01)

Package: app.api.routes
Purpose: Knowledge Management APIs
Version: v1.0.0 (Updated with Cohere + Qdrant integration)

Changes from v0.0.1:
- Added /documents/{id}/process endpoint for real indexing
- Implemented Qdrant integration for vector storage
- Using Cohere embed-multilingual-v3.0 for Vietnamese support
- Real RAG query with similarity search
"""

from fastapi import (
    APIRouter,
    HTTPException,
    Depends,
    Query,
    UploadFile,
    File,
    Form,
    Body,
)
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import List, Optional, Dict
from loguru import logger
from pathlib import Path
import asyncio
import os
import shutil
import tempfile
from datetime import datetime
from app.api.middleware.auth import get_admin_user
from app.api.middleware.subscription_guard import check_active_subscription
from app.config.settings import settings

from app.api.schemas.knowledge_schemas import (
    DocumentResponse,
    DocumentListResponse,
    DocumentDetailResponse,
    UploadDocumentResponse,
    UploadErrorResponse,
    ProcessDocumentRequest,
    ProcessDocumentResponse,
    KGQueryRequest,
    KGQueryResponse,
    KGQueryResultItem,
    QueryKnowledgeRequest,
    QueryKnowledgeResponse,
    RetrievedChunk,
    DeleteDocumentResponse,
    KnowledgeBaseStatusResponse,
    HybridQueryRequest,
    HybridQueryResponse,
    ImageSearchResult,
)
from app.db.postgres.models import KnowledgeDocument
from app.db.postgres.session import get_db

# Initialize router - no global auth, add individually per endpoint
router = APIRouter(prefix="/knowledge", tags=["Knowledge Base"])


# Storage configuration
def _is_directory_writable(path: Path) -> bool:
    """Check whether a directory can be created and written to."""
    try:
        path.mkdir(parents=True, exist_ok=True)
        with tempfile.NamedTemporaryFile(dir=path, prefix=".write_test_", delete=True):
            pass
        return True
    except (PermissionError, OSError) as exc:
        logger.warning(f"Storage path not writable: {path} ({exc})")
        return False


def get_storage_dir() -> Path:
    """Resolve a writable storage directory for uploaded knowledge files."""
    configured_dir = Path(settings.UPLOAD_DIR)
    candidate_dirs = [
        configured_dir,
        Path("/app/uploads/documents"),
        Path("uploads/documents"),
    ]

    seen = set()
    for candidate in candidate_dirs:
        normalized = str(candidate)
        if normalized in seen:
            continue
        seen.add(normalized)

        if _is_directory_writable(candidate):
            if candidate != configured_dir:
                logger.warning(
                    f"Upload dir '{configured_dir}' is not writable. Falling back to '{candidate}'."
                )
            return candidate.resolve()

    raise RuntimeError("Khong co thu muc luu tru nao co quyen ghi cho knowledge upload")


ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "md"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB - Hỗ trợ PDF lớn (300+ trang)


def get_rag_engine():
    """Lazy import RAG engine to avoid circular imports - Full LlamaIndex"""
    from app.core.rag.rag_engine import get_rag_engine as _get_rag_engine

    return _get_rag_engine()


# ===== UPLOAD DOCUMENT =====


@router.post(
    "/upload",
    response_model=UploadDocumentResponse,
    summary="[KB-01] Upload document for RAG",
    description="""
    Upload a document to the knowledge base.
    
    Supported formats: PDF, DOCX, TXT, MD
    Max file size: 50MB (phù hợp cho PDF 300+ trang)
    
    After upload, document needs to be processed to create vector embeddings.
    """,
)
async def upload_document(
    file: UploadFile = File(...),
    notes: Optional[str] = Form(None),
    uploaded_by: Optional[str] = Form("admin"),
    db: AsyncSession = Depends(get_db),
    _Subscription: bool = Depends(check_active_subscription),
):
    """
    Upload document to knowledge base

    Form Data:
        - file: Document file (PDF, DOCX, TXT, MD)
        - notes: Optional notes about the document
        - uploaded_by: Admin username
    """
    try:
        # Validate file extension
        filename = file.filename or "unknown"
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if extension not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type '{extension}' not allowed. Allowed: {list(ALLOWED_EXTENSIONS)}",
            )

        # Read file content
        content = await file.read()
        file_size = len(content)

        # Check file size
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE / 1024 / 1024}MB",
            )

        storage_dir = get_storage_dir()

        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{filename}"
        file_path = storage_dir / safe_filename

        # Save file
        with open(file_path, "wb") as f:
            f.write(content)

        # Create database record
        document = KnowledgeDocument(
            filename=filename,
            file_path=str(file_path),
            file_type=extension,
            file_size=file_size,
            processed=False,
            vector_count=0,
            uploaded_by=uploaded_by,
            notes=notes,
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)

        logger.info(f"Uploaded document: {filename} (ID: {document.id})")

        return UploadDocumentResponse(
            success=True,
            message=f"Tài liệu '{filename}' tải lên thành công",
            document_id=document.id,
            filename=filename,
            file_size=file_size,
            file_type=extension,
            status="pending",
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== PROCESS DOCUMENT (INDEX TO QDRANT) =====


@router.post(
    "/documents/{document_id}/process",
    response_model=ProcessDocumentResponse,
    summary="[KB-01] Process document for RAG",
    description="""
    Process uploaded document and create vector embeddings.

    This endpoint:
    1. Reads the document file
    2. Chunks the content using LlamaIndex
    3. Creates embeddings using Cohere embed-multilingual-v3.0
    4. Stores vectors in Qdrant Cloud

    After processing, the document can be queried via RAG.
    """,
)
async def process_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    _Subscription: bool = Depends(check_active_subscription),
):
    """
    Process document and index to Qdrant

    Path params:
        - document_id: ID of the uploaded document

    Returns:
        - chunks_created: Number of chunks indexed
        - processing_time_ms: Time taken to process
    """
    try:
        import time

        start_time = time.time()

        # Get document from database
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(
                status_code=404, detail=f"Không tìm thấy tài liệu {document_id}"
            )

        # Allow reprocessing if document has 0 vectors (failed previous attempt)
        if (
            document.processed
            and document.vector_count > 0
            and document.image_count > 0
        ):
            return ProcessDocumentResponse(
                success=True,
                message=f"Tài liệu '{document.filename}' đã được xử lý trước đó",
                document_id=document_id,
                chunks_created=document.vector_count,
                processing_time_ms=0,
            )

        if document.processed and (
            document.vector_count == 0 or document.image_count == 0
        ):
            logger.warning(
                f"Document {document_id} was marked processed with {document.vector_count} vectors, {document.image_count} images. Reprocessing for missing data..."
            )
            # Reset processed status for retry
            document.processed = False
            await db.commit()

        # Read file content
        file_path = document.file_path
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(
                status_code=404, detail=f"Không tìm thấy file tại {file_path}"
            )

        with open(file_path, "rb") as f:
            file_content = f.read()

        # Get RAG engine and index document
        rag = get_rag_engine()
        index_result = await rag.index_document(
            file_path=Path(file_path),
            filename=document.filename,
            document_id=document.id,
            metadata={
                "file_type": document.file_type,
                "uploaded_by": document.uploaded_by,
                "notes": document.notes,
            },
        )

        # Validate processing succeeded
        if index_result.text_chunks == 0:
            raise HTTPException(
                status_code=500,
                detail="Xử lý tài liệu thất bại: Không tạo được vectors. "
                "Nguyên nhân có thể: "
                "(1) Chưa cấu hình COHERE_API_KEY, "
                "(2) Chưa cấu hình QDRANT_URL, "
                "(3) API key không hợp lệ. "
                "Vui lòng kiểm tra cấu hình trong trang Knowledge.",
            )

        # Update document status (only if vectors were created successfully)
        document.processed = True
        document.vector_count = index_result.text_chunks
        document.image_count = index_result.image_vectors
        from datetime import timezone

        document.processed_at = datetime.now(timezone.utc)
        await db.commit()

        processing_time = int((time.time() - start_time) * 1000)

        logger.info(
            f"Processed document {document_id}: {index_result.text_chunks} text chunks, "
            f"{index_result.image_vectors} images in {processing_time}ms"
        )

        return ProcessDocumentResponse(
            success=True,
            message=f"Tài liệu '{document.filename}' xử lý thành công",
            document_id=document_id,
            chunks_created=index_result.text_chunks,
            images_indexed=index_result.image_vectors,
            processing_time_ms=processing_time,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing document {document_id}: {e}")

        # Cleanup on failure: Delete file and database record
        # Only cleanup if document was just uploaded (processed=False, vector_count=0)
        try:
            # Try to get document from database
            result = await db.execute(
                select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
            )
            doc = result.scalar_one_or_none()

            if doc and not doc.processed and doc.vector_count == 0:
                logger.warning(
                    f"Processing failed for new document {document_id}. Cleaning up..."
                )

                # Delete file from disk
                if doc.file_path and os.path.exists(doc.file_path):
                    os.remove(doc.file_path)
                    logger.info(f"Deleted file: {doc.file_path}")

                # Delete database record
                await db.delete(doc)
                await db.commit()
                logger.info(f"Deleted database record for document {document_id}")

        except Exception as cleanup_error:
            logger.error(f"Cleanup failed: {cleanup_error}")
            # Don't raise cleanup error, raise original error instead

        raise HTTPException(status_code=500, detail=str(e))


# ===== LIST DOCUMENTS =====


@router.get(
    "/documents",
    response_model=DocumentListResponse,
    summary="List all documents",
    description="Get all documents in knowledge base with processing status",
)
async def list_documents(
    processed: Optional[bool] = Query(None, description="Filter by processed status"),
    file_type: Optional[str] = Query(None, description="Filter by file type"),
    db: AsyncSession = Depends(get_db),
):
    """
    List all documents in knowledge base

    Query params:
        - processed: true/false
        - file_type: pdf/docx/txt/md
    """
    try:
        query = select(KnowledgeDocument)

        if processed is not None:
            query = query.where(KnowledgeDocument.processed == processed)

        if file_type:
            query = query.where(KnowledgeDocument.file_type == file_type)

        result = await db.execute(query)
        documents = result.scalars().all()

        # Count processed vs pending
        processed_count = sum(1 for d in documents if d.processed)
        pending_count = len(documents) - processed_count

        return DocumentListResponse(
            total=len(documents),
            processed_count=processed_count,
            pending_count=pending_count,
            documents=[DocumentResponse.model_validate(d) for d in documents],
        )

    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== GET DOCUMENT =====


@router.get(
    "/documents/{document_id}",
    response_model=DocumentDetailResponse,
    summary="Get document detail",
    description="Get document details with chunks preview",
)
async def get_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """
    Get document detail with:
    - Document metadata
    - First 5 chunks preview (if processed)
    """
    try:
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(
                status_code=404, detail=f"Không tìm thấy tài liệu {document_id}"
            )

        # Pending: Get chunks from Qdrant when implemented
        chunks_preview = []

        return DocumentDetailResponse(
            document=DocumentResponse.model_validate(document),
            chunks_preview=chunks_preview,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== DOWNLOAD DOCUMENT (SERVE FILE FOR PREVIEW) =====

# Content-type mapping
CONTENT_TYPE_MAP = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain; charset=utf-8",
    "md": "text/plain; charset=utf-8",
}


@router.get(
    "/documents/{document_id}/download",
    summary="Download/Preview document file",
    description="""
    Serve the original document file for preview in the frontend.
    
    Returns the raw file with appropriate content-type:
    - PDF → application/pdf (for PDF viewer)
    - DOCX → application/vnd.openxmlformats-officedocument.wordprocessingml.document
    - TXT/MD → text/plain (for text display)
    """,
)
async def download_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """
    Serve original document file for frontend preview

    Path params:
        - document_id: ID of the document

    Returns:
        FileResponse with the original file
    """
    try:
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(
                status_code=404, detail=f"Không tìm thấy tài liệu {document_id}"
            )

        file_path = document.file_path
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(
                status_code=404,
                detail=f"Không tìm thấy file trên đĩa cho tài liệu {document_id}",
            )

        # Determine content type
        content_type = CONTENT_TYPE_MAP.get(
            document.file_type or "", "application/octet-stream"
        )

        logger.info(
            f"Serving document {document_id}: {document.filename} ({content_type})"
        )

        return FileResponse(
            path=file_path,
            media_type=content_type,
            filename=document.filename,
            headers={
                "Content-Disposition": f'inline; filename="{document.filename}"',
                "Access-Control-Expose-Headers": "Content-Disposition",
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== DELETE DOCUMENT =====


@router.delete(
    "/documents/{document_id}",
    response_model=DeleteDocumentResponse,
    summary="Delete document",
    description="Delete document and its vectors from knowledge base",
)
async def delete_document(document_id: int, db: AsyncSession = Depends(get_db)):
    """
    Delete document:
    1. Remove file from storage
    2. Remove vectors from Qdrant (when implemented)
    3. Remove database record
    """
    try:
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            raise HTTPException(
                status_code=404, detail=f"Không tìm thấy tài liệu {document_id}"
            )

        filename = document.filename
        vector_count = document.vector_count
        file_path = document.file_path
        vectors_actually_deleted = 0

        # Delete vectors from Qdrant FIRST (most critical operation)
        # If this fails, we abort the entire delete operation
        if document.processed and vector_count > 0:
            try:
                rag = get_rag_engine()
                vectors_actually_deleted = await rag.delete_document(document_id)
                logger.info(
                    f"Deleted {vectors_actually_deleted} vectors from Qdrant for document {document_id}"
                )
            except Exception as e:
                # CRITICAL: If Qdrant delete fails, abort entire operation
                error_msg = f"Failed to delete vectors from Qdrant: {str(e)}. Aborting document deletion to prevent orphaned data."
                logger.error(error_msg)
                raise HTTPException(status_code=500, detail=error_msg)

        # Only proceed with file and DB deletion if Qdrant delete succeeded
        # Delete file from storage
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Deleted file: {file_path}")

        # Delete database record
        await db.delete(document)
        await db.commit()

        logger.info(f"Deleted document: {filename} (ID: {document_id})")

        return DeleteDocumentResponse(
            success=True,
            message=f"Tài liệu '{filename}' và {vectors_actually_deleted} vectors đã xóa thành công",
            document_id=document_id,
            filename=filename,
            vectors_deleted=vectors_actually_deleted,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== QUERY KNOWLEDGE BASE =====


@router.post(
    "/query",
    response_model=QueryKnowledgeResponse,
    summary="[KB-01] Test RAG retrieval",
    description="""
    Test RAG retrieval query using Cohere embeddings + Qdrant.

    Admin can test what chunks are retrieved for a given query.
    This helps verify that the knowledge base is working correctly.

    Uses:
    - Cohere embed-multilingual-v3.0 for query embedding
    - Qdrant Cloud for vector similarity search
    """,
)
async def query_knowledge(
    request: QueryKnowledgeRequest,
    db: AsyncSession = Depends(get_db),
    _Subscription: bool = Depends(check_active_subscription),
):
    """
    Test RAG retrieval with Qdrant + Cohere

    Body:
        {
            "query": "Trieu chung cho bi non?",
            "top_k": 5,
            "min_score": 0.5
        }

    Response includes:
        - Retrieved chunks with relevance scores
        - Source document info
        - Retrieval time
    """
    try:
        import time

        start_time = time.time()

        # Get RAG engine
        rag = get_rag_engine()

        # Query knowledge base using Qdrant
        results = await rag.query(
            query=request.query, top_k=request.top_k, min_score=request.min_score
        )

        # Convert to response format
        chunks = [
            RetrievedChunk(
                document_id=r.document_id,
                document_name=r.document_name,
                chunk_index=r.chunk_index,
                content=r.content,
                score=r.score,
                metadata={"source": r.document_name},
            )
            for r in results
        ]

        # If no results, provide helpful message
        if not chunks:
            # Check if there are any processed documents
            result = await db.execute(
                select(func.count(KnowledgeDocument.id)).where(
                    KnowledgeDocument.processed == True
                )
            )
            processed_count = result.scalar() or 0

            if processed_count == 0:
                chunks.append(
                    RetrievedChunk(
                        document_id=0,
                        document_name="system",
                        chunk_index=0,
                        content=f"Chua co document nao duoc processed. Vui long upload va process document truoc khi query. Query: {request.query}",
                        score=0.0,
                        metadata={"type": "info"},
                    )
                )
            else:
                chunks.append(
                    RetrievedChunk(
                        document_id=0,
                        document_name="system",
                        chunk_index=0,
                        content=f"Khong tim thay ket qua phu hop voi min_score={request.min_score}. Thu giam min_score xuong 0.0 de xem tat ca ket qua. Neu van khong co ket qua, co the do: (1) Vector dimension mismatch - can recreate collection, (2) Query khong lien quan den noi dung document. Query: {request.query}",
                        score=0.0,
                        metadata={
                            "type": "info",
                            "suggestion": "Try lowering min_score to 0.0",
                        },
                    )
                )

        retrieval_time = int((time.time() - start_time) * 1000)

        return QueryKnowledgeResponse(
            success=True,
            query=request.query,
            total_chunks=len(chunks),
            chunks=chunks,
            retrieval_time_ms=retrieval_time,
        )

    except Exception as e:
        logger.error(f"Error querying knowledge base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== HYBRID QUERY (TEXT + IMAGE) =====


@router.post(
    "/query-hybrid",
    response_model=HybridQueryResponse,
    summary="[KB-01] Hybrid search (text + image)",
    description="""
    Hybrid search using both text and image embeddings.
    
    Use this when:
    - Query contains both text and image URLs
    - Want to find similar cases by image
    - Combined text + image similarity search
    
    Requires:
    - JINA_API_KEY configured in Knowledge settings
    - PDF documents with extracted images
    """,
)
async def query_hybrid(request: HybridQueryRequest, db: AsyncSession = Depends(get_db)):
    """
    Hybrid query with text and/or image search

    Body:
        {
            "query": "triệu chứng ghẻ",
            "image_urls": ["https://example.com/pet_lesion.jpg"],
            "top_k": 5,
            "min_score": 0.5
        }
    """
    try:
        import time

        start_time = time.time()

        rag = get_rag_engine()
        result = await rag.query_with_images(
            query=request.query,
            image_urls=request.image_urls,
            top_k=request.top_k,
            min_score=request.min_score,
        )

        text_chunks = [
            RetrievedChunk(
                document_id=r.document_id,
                document_name=r.document_name,
                chunk_index=r.chunk_index,
                content=r.content,
                score=r.score,
                metadata={"source": r.document_name},
            )
            for r in result.get("text_results", [])
        ]

        image_results = [
            ImageSearchResult(
                document_id=r.get("document_id", 0),
                filename=r.get("filename", ""),
                image_id=r.get("image_id", ""),
                score=r.get("score", 0.0),
                payload=r.get("payload"),
            )
            for r in result.get("image_results", [])
        ]

        retrieval_time = int((time.time() - start_time) * 1000)

        return HybridQueryResponse(
            success=True,
            query=request.query,
            text_results=text_chunks,
            image_results=image_results,
            has_image_query=result.get("has_image_query", False),
            total_text_results=len(text_chunks),
            total_image_results=len(image_results),
            retrieval_time_ms=retrieval_time,
        )

    except Exception as e:
        logger.error(f"Error in hybrid query: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== RECREATE COLLECTION =====


@router.post(
    "/recreate-collection",
    summary="[Admin] Recreate Qdrant collection",
    description="""
    Manually delete and recreate the Qdrant collection with correct dimensions.

    Use this when:
    - Switching embedding models (OpenAI ↔ Cohere)
    - Fixing dimension mismatches
    - Resetting the knowledge base

    WARNING: This will delete ALL vectors. Documents in database will remain but need reprocessing.
    """,
)
async def recreate_collection(db: AsyncSession = Depends(get_db)):
    """
    Delete and recreate Qdrant collection

    Returns:
        Success message with collection info
    """
    try:
        from app.core.rag import get_rag_engine, COHERE_EMBED_DIMENSION

        # Use RAG engine to recreate collection (LlamaIndex handles Qdrant internally)
        rag = get_rag_engine()
        success = await rag.recreate_collection()

        if not success:
            raise HTTPException(status_code=500, detail="Failed to recreate collection")

        # Reset all documents to unprocessed
        result = await db.execute(select(KnowledgeDocument))
        documents = result.scalars().all()

        for doc in documents:
            doc.processed = False
            doc.vector_count = 0
            doc.processed_at = None

        await db.commit()

        # Get collection info after recreation
        status = await rag.get_status()

        logger.info(f"Recreated collection with dimension {COHERE_EMBED_DIMENSION}")
        logger.info(f"Reset {len(documents)} documents to unprocessed")

        return {
            "success": True,
            "message": f"Tái tạo collection thành công",
            "collection_name": status.get("collection_name"),
            "dimension": COHERE_EMBED_DIMENSION,
            "documents_reset": len(documents),
            "engine": "LlamaIndex",
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error recreating collection: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== DEBUG QDRANT =====


@router.get(
    "/debug/qdrant",
    summary="[Debug] Check Qdrant collection details",
    description="Debug endpoint to verify Qdrant collection configuration and contents",
)
async def debug_qdrant(db: AsyncSession = Depends(get_db)):
    """
    Debug Qdrant collection

    Returns detailed info about:
    - Collection existence
    - Vector dimensions
    - Number of points
    - Sample points
    """
    try:
        from app.core.rag import get_rag_engine, COHERE_EMBED_DIMENSION

        # Use RAG engine for debug info (LlamaIndex handles Qdrant internally)
        rag = get_rag_engine()
        debug_info = await rag.get_debug_info()

        # Check for error in debug_info
        if "error" in debug_info:
            return {
                "exists": False,
                "error": debug_info.get("error"),
                "collection_name": debug_info.get("collection_name"),
                "message": "Collection not accessible",
            }

        return {
            "exists": True,
            "collection_name": debug_info.get("collection_name"),
            "vectors_count": debug_info.get("vectors_count"),
            "status": debug_info.get("status"),
            "expected_dimension": COHERE_EMBED_DIMENSION,
            "sample_points": debug_info.get("sample_points", []),
            "engine": debug_info.get("engine"),
            "message": "Collection found and accessible",
        }

    except Exception as e:
        logger.error(f"Error debugging Qdrant: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ===== KNOWLEDGE BASE STATUS =====


@router.get(
    "/status",
    response_model=KnowledgeBaseStatusResponse,
    summary="Get knowledge base status",
    description="Overall status of the knowledge base including Qdrant info",
)
async def get_status(db: AsyncSession = Depends(get_db)):
    """
    Get overall knowledge base status

    Returns:
    - Document counts (total, processed, pending)
    - Vector counts from Qdrant
    - Storage size
    - Embedding model info
    """
    try:
        # Count documents
        total_result = await db.execute(select(func.count(KnowledgeDocument.id)))
        total_documents = total_result.scalar() or 0

        # Count processed
        processed_result = await db.execute(
            select(func.count(KnowledgeDocument.id)).where(
                KnowledgeDocument.processed == True
            )
        )
        processed_documents = processed_result.scalar() or 0

        # Sum text vectors from database
        vectors_result = await db.execute(
            select(func.sum(KnowledgeDocument.vector_count))
        )
        total_vectors = vectors_result.scalar() or 0

        # Sum image vectors from database
        image_vectors_result = await db.execute(
            select(func.sum(KnowledgeDocument.image_count))
        )
        total_image_vectors = image_vectors_result.scalar() or 0

        # Sum file sizes
        size_result = await db.execute(select(func.sum(KnowledgeDocument.file_size)))
        storage_size = size_result.scalar() or 0

        # Get last updated
        last_result = await db.execute(
            select(KnowledgeDocument.uploaded_at)
            .order_by(KnowledgeDocument.uploaded_at.desc())
            .limit(1)
        )
        last_updated = last_result.scalar()

        # Get Qdrant stats if available
        qdrant_info = {}
        try:
            rag = get_rag_engine()
            qdrant_info = await rag.get_status()
        except Exception as e:
            logger.warning(f"Could not get Qdrant status: {e}")
            qdrant_info = {"status": "unavailable", "error": str(e)}

        return KnowledgeBaseStatusResponse(
            total_documents=total_documents,
            processed_documents=processed_documents,
            pending_documents=total_documents - processed_documents,
            total_vectors=total_vectors,
            total_image_vectors=total_image_vectors,
            storage_size_bytes=storage_size,
            last_updated=last_updated,
            qdrant_info=qdrant_info,
        )

    except Exception as e:
        logger.error(f"Error getting knowledge base status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================
# KNOWLEDGE GRAPH (KG) ENDPOINTS  [KB-04]
# =============================================================


def get_kg_service():
    """Lazy import KnowledgeGraphService to avoid circular imports."""
    from app.core.rag.knowledge_graph import get_knowledge_graph_service

    return get_knowledge_graph_service()


@router.post(
    "/build-kg",
    summary="[KB-04] Build Knowledge Graph từ documents đã xử lý",
    description="""
    Trích xuất triplets (subject, predicate, object) từ tài liệu đã processed
    và xây dựng Knowledge Graph.

    Flow:
    1. Đọc tất cả documents đã processed từ PostgreSQL
    2. Load nội dung file -> tạo LlamaIndex Document objects
    3. LLM extract triplets từ mỗi chunk
    4. Lưu vào SimpleGraphStore (persist to disk)

    Lưu ý: Quá trình này có thể mất vài phút tùy số lượng tài liệu.
    """,
    dependencies=[Depends(get_admin_user)],
)
async def build_knowledge_graph(
    document_ids: Optional[List[int]] = Query(
        None, description="IDs tài liệu cụ thể. Để trống = tất cả đã processed."
    ),
    max_triplets: int = Query(
        default=200,
        ge=1,
        le=1000,
        description="Số triplets tối đa tổng cộng sau deduplication",
    ),
    db: AsyncSession = Depends(get_db),
):
    """Build/extend Knowledge Graph từ processed documents."""
    import time

    start_time = time.time()

    try:
        # Query documents
        query = select(KnowledgeDocument).where(KnowledgeDocument.processed == True)
        if document_ids:
            query = query.where(KnowledgeDocument.id.in_(document_ids))

        result = await db.execute(query)
        documents = result.scalars().all()

        # Check if there are documents but none are processed
        all_docs_query = select(KnowledgeDocument)
        all_result = await db.execute(all_docs_query)
        all_docs = all_result.scalars().all()

        if not documents:
            if all_docs:
                # Documents exist but none are processed
                processed_count = sum(1 for d in all_docs if d.processed)
                raise HTTPException(
                    status_code=400,
                    detail=f"Tìm thấy {len(all_docs)} tài liệu nhưng không có tài liệu nào được xử lý (processed). "
                    f"Để xử lý tài liệu, bạn cần: "
                    f"(1) Cấu hình COHERE_API_KEY và QDRANT_URL trong trang Knowledge, "
                    f"(2) Upload lại tài liệu (quá trình xử lý sẽ tự động chạy).",
                )
            else:
                raise HTTPException(
                    status_code=404,
                    detail="Không tìm thấy tài liệu nào trong hệ thống. Vui lòng upload tài liệu trước.",
                )

        # Read file contents and create LlamaIndex Documents
        from llama_index.core import Document as LlamaDocument
        import unicodedata

        def _normalize_text(text: str) -> str:
            text = unicodedata.normalize("NFC", text or "")
            cleaned = []
            for ch in text:
                cat = unicodedata.category(ch)
                if cat.startswith("C") and ch not in ("\n", "\r", "\t"):
                    continue
                cleaned.append(ch)
            return "".join(cleaned).strip()

        def _extract_text_from_file(path: Path, file_type: Optional[str]) -> str:
            ft = (file_type or "").lower().strip()

            if ft in ("txt", "md"):
                return path.read_text(encoding="utf-8", errors="replace")

            if ft == "docx":
                from docx import Document as DocxDocument

                docx = DocxDocument(str(path))
                return "\n".join(p.text for p in docx.paragraphs if p.text)

            if ft == "pdf":
                text_parts: List[str] = []
                # Prefer PyMuPDF if available
                try:
                    import fitz  # PyMuPDF

                    with fitz.open(str(path)) as pdf:
                        for page in pdf:
                            t = page.get_text("text") or ""
                            if t.strip():
                                text_parts.append(t)
                    return "\n".join(text_parts)
                except Exception:
                    from PyPDF2 import PdfReader

                    reader = PdfReader(str(path))
                    for page in reader.pages:
                        t = page.extract_text() or ""
                        if t.strip():
                            text_parts.append(t)
                    return "\n".join(text_parts)

            return path.read_text(encoding="utf-8", errors="replace")

        # Get storage directory for resolving relative paths
        storage_dir = get_storage_dir()
        logger.info(f"Build KG using storage_dir: {storage_dir}")

        llama_docs: List[LlamaDocument] = []
        skipped: List[int] = []
        skipped_reasons: Dict[int, str] = {}
        for doc in documents:
            # Resolve file path - try multiple approaches
            doc_path = None

            # 1. Try as absolute path first
            if Path(doc.file_path).is_absolute():
                doc_path = Path(doc.file_path)

            # 2. Try relative to storage_dir
            if not doc_path or not doc_path.exists():
                doc_path = storage_dir / doc.file_path

            # 3. Try just the filename in storage_dir
            if not doc_path or not doc_path.exists():
                doc_path = storage_dir / Path(doc.file_path).name

            logger.info(
                f"Document {doc.id} ({doc.filename}): stored_path='{doc.file_path}', resolved='{doc_path}', exists={doc_path.exists() if doc_path else False}"
            )

            if not doc_path or not doc_path.exists():
                logger.warning(
                    f"Document {doc.id} file not found after trying all paths"
                )
                skipped.append(doc.id)
                skipped_reasons[doc.id] = "Không tìm thấy file trên ổ đĩa"
                continue
            try:
                # Use asyncio.to_thread for blocking file IO
                raw_text = await asyncio.to_thread(
                    _extract_text_from_file, doc_path, doc.file_type
                )
                text = _normalize_text(raw_text)

                # PDF scan/image-only thường gần như không có text -> KG không thể extract
                if len(text) < 200:
                    logger.warning(
                        f"Document {doc.id} has too little text for KG extraction (len={len(text)})."
                    )
                    skipped.append(doc.id)
                    skipped_reasons[doc.id] = (
                        "Tài liệu quá ít chữ (có thể là PDF dạng hình ảnh). "
                        "Vui lòng dùng tài liệu có text hoặc bổ sung OCR."
                    )
                    continue

                logger.info(
                    f"Document {doc.id} loaded successfully for KG, text length: {len(text)}"
                )
                llama_docs.append(
                    LlamaDocument(
                        text=text,
                        metadata={
                            "document_id": doc.id,
                            "filename": doc.filename,
                            "file_type": doc.file_type,
                        },
                    )
                )
            except Exception as e:
                logger.warning(f"Could not read document {doc.id}: {e}")
                skipped.append(doc.id)
                skipped_reasons[doc.id] = f"Lỗi đọc nội dung: {e}"

        if not llama_docs:
            raise HTTPException(
                status_code=400,
                detail=f"Không thể đọc nội dung từ tài liệu. "
                f"Đã kiểm tra {len(documents)} tài liệu, {len(skipped)} bị bỏ qua do lỗi đọc file. "
                f"Vui lòng kiểm tra: (1) File có tồn tại trong thư mục uploads/documents? (2) Thư mục lưu trữ có đúng không?",
            )

        # Build KG
        kg = get_kg_service()
        triplet_count = await kg.build_from_documents(
            llama_docs, max_triplets_per_chunk=max_triplets
        )

        processing_time = int((time.time() - start_time) * 1000)

        return {
            "success": True,
            "message": f"Knowledge Graph đã xây dựng thành công",
            "documents_processed": len(llama_docs),
            "documents_skipped": skipped,
            "documents_skipped_reasons": skipped_reasons,
            "triplets_extracted": triplet_count,
            "processing_time_ms": processing_time,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error building Knowledge Graph: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================
# PUBLIC KG ENDPOINTS (No Auth Required)
# =============================================================


@router.get(
    "/kg-stats",
    summary="[KB-04] Thống kê Knowledge Graph",
    description="Thông tin chi tiết về Knowledge Graph: số triplets, entities, relation types.",
)
async def get_kg_stats():
    """Get Knowledge Graph statistics."""
    try:
        kg = get_kg_service()
        stats = await kg.get_graph_stats()
        return {"success": True, **stats}
    except Exception as e:
        logger.error(f"Error getting KG stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/kg-visualize",
    summary="[KB-05] Lấy dữ liệu để visualize Knowledge Graph",
    description="Trả về nodes và edges dạng JSON cho D3.js visualization.",
)
async def get_kg_visualize():
    """Get Knowledge Graph data for visualization."""
    try:
        kg = get_kg_service()
        data = await kg.get_graph_visualization_data()
        return {"success": True, **data}
    except Exception as e:
        logger.error(f"Error getting KG visualization data: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/kg-query",
    response_model=KGQueryResponse,
    summary="[KB-05] Truy vấn Knowledge Graph",
    description="Truy vấn tri thức có cấu trúc từ đồ thị tri thức (triplets).",
)
async def query_knowledge_graph(request: KGQueryRequest):
    """
    Query the Knowledge Graph directly.

    Returns structured triplets related to the query.
    """
    try:
        kg = get_kg_service()
        results = await kg.query_graph(request.query, top_k=request.top_k)

        logger.info(f"[KG Query API] Got {len(results)} results")

        formatted_results = []

        # Now query_graph returns structured triplets in triplets_used
        for r in results:
            logger.info(
                f"[KG Query API] Result: content={r.content[:50] if r.content else 'None'}..., triplets_used={len(r.triplets_used) if r.triplets_used else 0}"
            )

            if r.triplets_used:
                # Use actual triplets from the response
                for t in r.triplets_used:
                    logger.info(f"[KG Query API] Triplet type: {type(t)}, value: {t}")
                    if isinstance(t, dict):
                        formatted_results.append(
                            KGQueryResultItem(
                                subject=t.get("subject", ""),
                                predicate=t.get("predicate", ""),
                                object=t.get("object", ""),
                                score=r.score,
                                source_nodes=r.source_nodes,
                            )
                        )
                    else:
                        # Fallback for non-dict format
                        logger.warning(f"[KG Query API] Triplet is not a dict: {t}")

        return KGQueryResponse(
            success=True,
            query=request.query,
            results=formatted_results,
            message=f"Tìm thấy {len(formatted_results)} triplets từ Knowledge Graph",
        )
    except Exception as e:
        logger.error(f"Error querying KG: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================
# CASE MEMORY ENDPOINTS  [KB-05]
# =============================================================


def get_cm_service():
    """Lazy import CaseMemoryService to avoid circular imports."""
    from app.core.rag.case_memory import get_case_memory_service

    return get_case_memory_service()


@router.get(
    "/case-memory/stats",
    summary="[KB-05] Thống kê Case Memory",
    description="Thông tin về Qdrant collection `petties_case_memory`: số cases, status.",
)
async def get_case_memory_stats():
    """Get Case Memory collection statistics."""
    try:
        cm = get_cm_service()
        stats = await cm.get_stats()
        return {"success": True, **stats}
    except Exception as e:
        logger.error(f"Error getting Case Memory stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/case-memory/prune",
    summary="[KB-05] Dọn dẹp Case Memory",
    description="""
    Xóa các cases cũ không có feedback (feedback_count = 0) để giữ collection sạch.

    Chỉ xóa cases cũ hơn `older_than_days` ngày.
    """,
)
async def prune_case_memory(
    older_than_days: int = Query(
        default=90, ge=1, le=365, description="Chỉ xóa cases cũ hơn X ngày"
    ),
    max_feedback_below: int = Query(
        default=0, ge=0, le=5, description="Xóa cases có feedback_count <= X"
    ),
):
    """Prune low-score / stale cases from Case Memory."""
    try:
        cm = get_cm_service()
        pruned_count = await cm.prune_low_score_cases(
            max_feedback_below=max_feedback_below,
            older_than_days=older_than_days,
        )
        return {
            "success": True,
            "message": f"Đã xóa {pruned_count} cases không có feedback",
            "pruned_count": pruned_count,
            "criteria": {
                "max_feedback_below": max_feedback_below,
                "older_than_days": older_than_days,
            },
        }
    except Exception as e:
        logger.error(f"Error pruning Case Memory: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/case-memory",
    summary="[KB-05] Danh sách Cases",
    description="Lấy danh sách cases với pagination và filters (species, diagnosis, query)",
)
async def list_case_memory(
    query: Optional[str] = Query(
        default=None, description="Tìm kiếm trong nội dung case"
    ),
    species: Optional[str] = Query(
        default=None, description="Lọc theo loài (dog, cat, other)"
    ),
    diagnosis: Optional[str] = Query(
        default=None, description="Lọc theo từ khóa chẩn đoán"
    ),
    page: int = Query(default=1, ge=1, description="Số trang"),
    page_size: int = Query(default=20, ge=1, le=100, description="Số items mỗi trang"),
    _: dict = Depends(get_admin_user),
):
    """List cases with pagination and filters."""
    try:
        cm = get_cm_service()
        result = await cm.list_cases(
            query=query,
            species=species,
            diagnosis=diagnosis,
            page=page,
            page_size=page_size,
        )
        return {"success": True, **result}
    except Exception as e:
        logger.error(f"Error listing Case Memory: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get(
    "/case-memory/{case_id}",
    summary="[KB-05] Chi tiết Case",
    description="Lấy chi tiết một case theo ID",
)
async def get_case_memory(
    case_id: str,
    _: dict = Depends(get_admin_user),
):
    """Get case detail by ID."""
    try:
        cm = get_cm_service()
        case = await cm.get_case(case_id)
        if case is None:
            raise HTTPException(status_code=404, detail="Không tìm thấy case")
        return {"success": True, "case": case}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting Case Memory {case_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.patch(
    "/case-memory/{case_id}",
    summary="[KB-05] Cập nhật Case",
    description="Cập nhật metadata của một case (diagnosis, symptoms)",
)
async def update_case_memory(
    case_id: str,
    diagnosis: Optional[str] = Body(default=None, description="Chẩn đoán mới"),
    symptoms: Optional[List[str]] = Body(
        default=None, description="Danh sách triệu chứng mới"
    ),
    _: dict = Depends(get_admin_user),
):
    """Update case metadata."""
    try:
        cm = get_cm_service()
        success = await cm.update_case(
            case_id=case_id,
            diagnosis=diagnosis,
            symptoms=symptoms,
        )
        if not success:
            raise HTTPException(
                status_code=404, detail="Không tìm thấy case hoặc cập nhật thất bại"
            )
        return {"success": True, "message": "Cập nhật case thành công"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating Case Memory {case_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete(
    "/case-memory/{case_id}",
    summary="[KB-05] Xóa Case",
    description="Xóa một case khỏi Case Memory",
)
async def delete_case_memory(
    case_id: str,
    _: dict = Depends(get_admin_user),
):
    """Delete a case from Case Memory."""
    try:
        cm = get_cm_service()
        success = await cm.delete_case(case_id)
        if not success:
            raise HTTPException(
                status_code=404, detail="Không tìm thấy case hoặc xóa thất bại"
            )
        return {"success": True, "message": "Xóa case thành công"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting Case Memory {case_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post(
    "/case-memory/sync-emr-confirmed",
    summary="[KB-05] Đồng bộ EMR confirmed vào Case Memory",
    description=(
        "Gọi Spring internal endpoint để lấy EMR confirmed và upsert vào Case Memory. "
        "Chỉ dành cho admin hoặc job vận hành nội bộ."
    ),
)
async def sync_emr_confirmed_into_case_memory(
    limit: int = Query(default=50, ge=1, le=200, description="Số EMR tối đa mỗi batch"),
    cursor: Optional[str] = Query(default=None, description="Cursor của batch trước"),
    updated_from: Optional[str] = Query(
        default=None, description="ISO datetime bắt đầu"
    ),
    updated_to: Optional[str] = Query(
        default=None, description="ISO datetime kết thúc"
    ),
    _: dict = Depends(get_admin_user),
):
    """Manually sync confirmed EMR records into Case Memory."""
    raise HTTPException(
        status_code=410,
        detail="Endpoint này đã ngưng sử dụng. Spring Boot sẽ push trực tiếp EMR sang AI service.",
    )
